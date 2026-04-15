from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
import os
import requests
import io
import json
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from werkzeug.utils import secure_filename
import tempfile

from docx import Document
from fpdf import FPDF
import matplotlib
matplotlib.use('Agg') 
import matplotlib.pyplot as plt

app = Flask(__name__)

# --- CONFIGURATION ---
API_KEY = 'helloworld' 
# ADDED .docx TO ALLOWED EXTENSIONS
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf', 'csv', 'xlsx', 'xls', 'json', 'txt', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_file(filepath, filename):
    ext = filename.rsplit('.', 1)[1].lower()
    df = None

    try:
        if ext in ['xlsx', 'xls']:
            df = pd.read_excel(filepath, header=None)
        elif ext == 'csv' or ext == 'txt':
            df = pd.read_csv(filepath, header=None, on_bad_lines='skip', sep=None, engine='python')
        elif ext == 'json':
            df = pd.read_json(filepath)
            
        # --- NEW: WORD DOCUMENT UPLOAD SUPPORT ---
        elif ext == 'docx':
            doc = Document(filepath)
            data = []
            for table in doc.tables:
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
            if data:
                df = pd.DataFrame(data)
            else:
                return None, "No tables found in the Word document."
                
        elif ext in ['png', 'jpg', 'jpeg', 'pdf']:
            url = "https://api.ocr.space/parse/image"
            payload = {'apikey': API_KEY, 'isTable': True, 'scale': True}
            with open(filepath, 'rb') as f:
                response = requests.post(url, files={'file': f}, data=payload)
            
            result = response.json()
            if result.get('IsErroredOnProcessing'):
                return None, result.get('ErrorMessage')[0]
                
            text = result['ParsedResults'][0]['ParsedText']
            lines = [line.split('\t') for line in text.split('\r\n') if line.strip()]
            df = pd.DataFrame(lines)
        
        # --- CLEANING LOGIC ---
        if df is not None and not df.empty:
            df.dropna(how='all', inplace=True) 
            df.dropna(axis=1, how='all', inplace=True)
            valid_rows = df.dropna(thresh=min(3, len(df.columns))) 
            
            if not valid_rows.empty:
                header_idx = valid_rows.index[0]
                new_header = df.loc[header_idx] 
                df = df.loc[header_idx + 1:] 
                df.columns = new_header 
                
            df.reset_index(drop=True, inplace=True)
            df.columns = [str(col).strip().upper() for col in df.columns]
            df = df.fillna("") 
            df = df.replace({r'[\x00-\x08\x0b-\x0c\x0e-\x1f]': ''}, regex=True)
            return df, None
            
        return None, "Could not extract structured data."
    except Exception as e:
        return None, str(e)

# --- WEB ROUTES ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'})
        
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        file.save(filepath)
        
        df, error = process_file(filepath, filename)
        os.remove(filepath)
        
        if error:
            return jsonify({'error': error})
            
        table_html = df.to_html(classes='min-w-full text-left text-sm font-light border-collapse', index=False, border=0)
        return jsonify({'message': 'Success', 'table': table_html, 'json_data': df.to_dict('records')})

    return jsonify({'error': 'Invalid file type'})

@app.route('/export', methods=['POST'])
def export_data():
    data = request.json.get('data')
    format_type = request.json.get('format')
    df = pd.DataFrame(data)
    output = io.BytesIO()
    
    mimetype = ''
    filename = ''

    if format_type == 'excel':
        df.to_excel(output, index=False, engine='openpyxl')
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        filename = 'Cleaned_Data.xlsx'
    
    elif format_type == 'csv':
        df.to_csv(output, index=False)
        mimetype = 'text/csv'
        filename = 'Cleaned_Data.csv'
        
    elif format_type == 'json':
        df.to_json(output, orient='records', indent=4)
        mimetype = 'application/json'
        filename = 'Cleaned_Data.json'
        
    elif format_type == 'txt':
        df.to_csv(output, index=False, sep='\t')
        mimetype = 'text/plain'
        filename = 'Cleaned_Data.txt'
        
    elif format_type == 'docx':
        doc = Document()
        doc.add_heading('Data Export', 0)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(df.columns): 
            table.rows[0].cells[i].text = str(col)
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row): 
                row_cells[i].text = str(item)
        doc.save(output)
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = 'Cleaned_Data.docx'
        
    elif format_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=9)
        col_width = 190 / (len(df.columns) or 1)
        
        # FIX: Sanitize text to prevent PDF from crashing the server on weird symbols
        def sanitize(text):
            return str(text)[:20].encode('latin-1', 'replace').decode('latin-1')
            
        for col in df.columns:
            pdf.cell(col_width, 10, sanitize(col), border=1)
        pdf.ln()
        for _, row in df.iterrows():
            for item in row:
                pdf.cell(col_width, 10, sanitize(item), border=1)
            pdf.ln()
        pdf_bytes = pdf.output(dest='S').encode('latin1')
        output.write(pdf_bytes)
        mimetype = 'application/pdf'
        filename = 'Cleaned_Data.pdf'
        
    elif format_type == 'image':
        # FIX: Dynamic High-Res Image Sizing
        cols = len(df.columns)
        rows = len(df) + 1
        fig, ax = plt.subplots(figsize=(max(10, cols * 2), max(5, rows * 0.6)))
        ax.axis('tight')
        ax.axis('off')
        
        table = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(12)
        table.scale(1, 2) # Adds padding to rows so text isn't blurry or cramped
        
        plt.savefig(output, format='png', bbox_inches='tight', dpi=400) # Upgraded to 400 DPI
        plt.close('all') # CRITICAL: Frees server memory so it doesn't crash on the next download
        
        mimetype = 'image/png'
        filename = 'Cleaned_Data.png'

    output.seek(0)
    return send_file(output, mimetype=mimetype, as_attachment=True, download_name=filename)

@app.route('/append_export', methods=['POST'])
def append_export():
    if 'existing_file' not in request.files:
        return jsonify({'error': 'No existing file provided'})
        
    existing_file = request.files['existing_file']
    new_data_json = request.form.get('new_data')
    
    if not new_data_json or existing_file.filename == '':
        return jsonify({'error': 'Missing data or file'})
        
    ext = existing_file.filename.rsplit('.', 1)[1].lower()
    df_new = pd.DataFrame(json.loads(new_data_json))
    output = io.BytesIO()
    
    try:
        if ext in ['xlsx', 'xls']:
            wb = load_workbook(existing_file)
            ws = wb.active
            for r in dataframe_to_rows(df_new, index=False, header=False): 
                ws.append(r)
            wb.save(output)
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            
        elif ext == 'csv':
            existing_df = pd.read_csv(existing_file)
            combined_df = pd.concat([existing_df, df_new], ignore_index=True)
            combined_df.to_csv(output, index=False)
            mimetype = 'text/csv'
        else:
            return jsonify({'error': 'Can only append to Excel or CSV files.'})
            
        output.seek(0)
        download_name = 'Updated_' + secure_filename(existing_file.filename)
        return send_file(output, mimetype=mimetype, as_attachment=True, download_name=download_name)
    except Exception as e:
        return jsonify({'error': f'Failed to append: {str(e)}'})

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
